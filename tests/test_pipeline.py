from __future__ import annotations

from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from PIL import Image, ImageDraw

from ocr_app.bijoy import convert_unicode_to_bijoy
from ocr_app.docx_builder import export_all
from ocr_app.docx_builder import build_docx
from ocr_app.docx_builder import _content_width_weights
from ocr_app.docx_builder import _crop_artifact
from ocr_app.docx_builder import _detect_table_grid_geometry
from ocr_app.docx_builder import _apply_table_geometry
from ocr_app.docx_builder import _prepare_faithful_blocks
from ocr_app.main import app
from ocr_app.models import BoundingBox, DocumentLayout, LayoutBlock, PageLayout, TableBlock, TableCell
from ocr_app.ocr import extract_layout
from ocr_app.ocr import _expand_compressed_layout, _normalise_bbox, _restore_colon_separator_column
from ocr_app.pdf_render import render_pdf_pages
from ocr_app.settings import clear_api_key
from ocr_app.storage import DEMO_DIRECT_MATTER, create_job_dir, save_demo_pdf, write_layout


def test_demo_reference_pipeline_creates_editable_docx() -> None:
    job_dir = create_job_dir()
    pdf = save_demo_pdf(job_dir, DEMO_DIRECT_MATTER / "04. Railway.pdf")
    pages = render_pdf_pages(pdf, job_dir / "pages")
    layout = extract_layout(pdf, pages, model="gemini-3-flash-preview", force_demo=True)
    write_layout(job_dir, layout)
    result = export_all(layout, job_dir / "exports")

    faithful = Path(result["faithful"])
    column_ready = Path(result["column_ready"])
    assert faithful.exists()
    assert Path(result["editable_faithful"]).exists()
    assert column_ready.exists()

    doc = Document(str(column_ready))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Bangladesh Railway" in text
    assert any(table.rows for table in doc.tables)


def test_api_demo_convert_and_export() -> None:
    client = TestClient(app)
    response = client.post(
        "/api/jobs",
        data={
            "demo_pdf": "05. Rajshahi.pdf",
            "model": "gemini-3-flash-preview",
            "force_demo": "true",
        },
    )
    assert response.status_code == 200
    payload = response.json()
    job_id = payload["job_id"]
    assert payload["layout"]["pages"]

    export_response = client.post(f"/api/jobs/{job_id}/export")
    assert export_response.status_code == 200
    export_payload = export_response.json()
    assert Path(export_payload["faithful"]).exists()
    assert Path(export_payload["column_ready"]).exists()


def test_api_key_settings_save_and_clear(monkeypatch, tmp_path) -> None:
    monkeypatch.delenv("GEMINI_API_KEY", raising=False)
    monkeypatch.setattr("ocr_app.settings.ENV_LOCAL_PATH", tmp_path / ".env.local")
    clear_api_key()
    client = TestClient(app)
    initial = client.get("/api/settings")
    assert initial.status_code == 200
    assert initial.json()["has_gemini_api_key"] is False

    saved = client.post("/api/settings/gemini-key", json={"api_key": "AIzaSyExampleLocalKey123456"})
    assert saved.status_code == 200
    assert saved.json()["has_gemini_api_key"] is True
    assert "..." in saved.json()["masked_gemini_api_key"]

    cleared = client.delete("/api/settings/gemini-key")
    assert cleared.status_code == 200
    assert cleared.json()["has_gemini_api_key"] is False


def test_numbered_form_table_restores_missing_colon_column() -> None:
    table = TableBlock(
        row_count=3,
        col_count=3,
        cells=[
            TableCell(row=0, col=0, text="১"),
            TableCell(row=0, col=1, text="কর্তৃপক্ষ"),
            TableCell(row=0, col=2, text="জেলা মৎস্য কর্মকর্তা"),
            TableCell(row=1, col=0, text="২"),
            TableCell(row=1, col=1, text="যে কাজের জন্য দরপত্র"),
            TableCell(row=1, col=2, text="দরপত্র প্রস্তাব অনুযায়ী"),
            TableCell(row=2, col=0, text="৩"),
            TableCell(row=2, col=1, text="দরপত্র সূত্র ও তারিখ"),
            TableCell(row=2, col=2, text="স্মারক নং- ৩৩.০২"),
        ],
    )
    restored = _restore_colon_separator_column(table)
    assert restored.col_count == 4
    assert [cell.text for cell in restored.cells if cell.row == 0 and cell.col in {2, 3}] == [":", "জেলা মৎস্য কর্মকর্তা"]


def test_unicode_to_bijoy_converter_outputs_ansi_text() -> None:
    converted, warnings = convert_unicode_to_bijoy("\u0986\u09ae\u09be\u09b0 \u09ac\u09be\u0982\u09b2\u09be\u0964")
    assert warnings == []
    assert converted == "Avgvi evsjv|"


def test_bijoy_export_keeps_english_runs_times_new_roman(tmp_path) -> None:
    image_path = tmp_path / "page.png"
    image_path.write_bytes(
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde"
        b"\x00\x00\x00\x0cIDATx\x9cc\xf8\xff\xff?\x00\x05\xfe\x02\xfeA\xbd\xb6\xef\x00\x00\x00\x00IEND\xaeB`\x82"
    )
    layout = DocumentLayout(
        source_pdf="sample.pdf",
        pages=[
            PageLayout(
                page_index=0,
                width_px=1,
                height_px=1,
                image_path=str(image_path),
                blocks=[
                    LayoutBlock(
                        id="mixed",
                        type="paragraph",
                        text="\u0986\u09ae\u09be\u09b0 English 2026",
                        bbox=BoundingBox(x=0.1, y=0.1, w=0.8, h=0.1),
                        confidence=1,
                    )
                ],
            )
        ],
    )
    result = export_all(layout, tmp_path / "exports")
    doc = Document(result["bijoy"])
    runs = [(run.text, run.font.name) for run in doc.paragraphs[0].runs if run.text]
    assert ("Avgvi", "SutonnyMJ") in runs
    assert (" English 2026", "Times New Roman") in runs


def test_four_column_form_table_uses_narrow_serial_and_colon_columns() -> None:
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=1, cols=4)
    block = LayoutBlock(
        id="table",
        type="table",
        bbox=BoundingBox(x=0.06, y=0.2, w=0.88, h=0.52),
        confidence=1,
    )
    table_block = TableBlock(
        row_count=2,
        col_count=4,
        cells=[
            TableCell(row=0, col=0, text="1"),
            TableCell(row=0, col=1, text="Label"),
            TableCell(row=0, col=2, text=":"),
            TableCell(row=0, col=3, text="Long value text"),
            TableCell(row=1, col=0, text="2"),
            TableCell(row=1, col=1, text="Label"),
            TableCell(row=1, col=2, text=":"),
            TableCell(row=1, col=3, text="Long value text"),
        ],
    )
    _apply_table_geometry(table, block, table_block=table_block, cols=4, compact=True)
    widths = [int(cell._tc.tcPr.tcW.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w")) for cell in table.rows[0].cells]
    assert widths[0] < widths[1]
    assert widths[2] < widths[0]
    assert widths[3] > widths[1]
    tblp = table._tbl.tblPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblpPr")
    assert tblp is not None
    assert tblp.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}horzAnchor") == "page"
    assert tblp.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vertAnchor") == "page"


def test_generic_tender_table_does_not_use_colon_form_widths() -> None:
    table_block = TableBlock(
        row_count=2,
        col_count=4,
        cells=[
            TableCell(row=0, col=0, text="Tender ID No."),
            TableCell(row=0, col=1, text="Package No"),
            TableCell(row=0, col=2, text="Description"),
            TableCell(row=0, col=3, text="Publication date\nClosing Date"),
            TableCell(row=1, col=0, text="1268623"),
            TableCell(row=1, col=1, text="Dmpi/Asset/2025-26/Re-tender/Gd-07"),
            TableCell(row=1, col=2, text="Goods Supply and Installation of Office Equipment"),
            TableCell(row=1, col=3, text="04-May-2026 16:00:00\n12-May-2026 11:00:00"),
        ],
    )
    weights = _content_width_weights(table_block, cols=4)
    assert weights[2] > 0.20
    assert weights[2] > weights[0]
    assert weights[3] > 0.16


def test_faithful_blocks_snap_memo_and_date_to_same_row() -> None:
    blocks = [
        LayoutBlock(
            id="memo",
            type="paragraph",
            text="Memo No. 59.14.8100.028.59.008.26-1505",
            bbox=BoundingBox(x=0.169, y=0.167, w=0.321, h=0.017),
            confidence=1,
            alignment="left",
        ),
        LayoutBlock(
            id="date",
            type="paragraph",
            text="Date:03-05-2026",
            bbox=BoundingBox(x=0.776, y=0.162, w=0.13, h=0.018),
            confidence=1,
            alignment="right",
        ),
    ]
    prepared = {block.id: block for block in _prepare_faithful_blocks(blocks)}
    assert prepared["memo"].bbox.y == prepared["date"].bbox.y


def test_signature_artifact_crop_expands_to_visible_ink(tmp_path) -> None:
    image_path = tmp_path / "page.png"
    image = Image.new("RGB", (500, 500), "white")
    draw = ImageDraw.Draw(image)
    draw.line((330, 330, 450, 305), fill="black", width=5)
    image.save(image_path)
    block = LayoutBlock(
        id="sig",
        type="artifact",
        artifact_type="signature",
        bbox=BoundingBox(x=0.70, y=0.61, w=0.04, h=0.025),
        confidence=1,
    )
    crop = _crop_artifact(block, image_path, tmp_path, 0)
    assert crop is not None
    cropped = Image.open(crop).convert("L")
    assert sum(1 for value in cropped.getdata() if value < 185) > 50


def test_column_export_uses_detected_signature_artifact_not_inferred_crop(tmp_path) -> None:
    image_path = tmp_path / "page.png"
    image = Image.new("RGB", (800, 1000), "white")
    draw = ImageDraw.Draw(image)
    draw.line((525, 810, 670, 775), fill="black", width=5)
    image.save(image_path)
    layout = DocumentLayout(
        source_pdf="sample.pdf",
        pages=[
            PageLayout(
                page_index=0,
                width_px=800,
                height_px=1000,
                image_path=str(image_path),
                blocks=[
                    LayoutBlock(
                        id="left_signatory",
                        type="paragraph",
                        text="Left officer\nDesignation\nOffice",
                        bbox=BoundingBox(x=0.18, y=0.78, w=0.28, h=0.08),
                        confidence=1,
                    ),
                    LayoutBlock(
                        id="right_signatory",
                        type="paragraph",
                        text="Right officer\nDesignation\nOffice",
                        bbox=BoundingBox(x=0.58, y=0.78, w=0.28, h=0.08),
                        confidence=1,
                    ),
                    LayoutBlock(
                        id="sig_artifact",
                        type="artifact",
                        artifact_type="signature",
                        bbox=BoundingBox(x=0.66, y=0.79, w=0.05, h=0.025),
                        confidence=1,
                    ),
                ],
            )
        ],
    )
    build_docx(layout, tmp_path / "column.docx", mode="column", bijoy=False)
    crop_dir = tmp_path / "artifact-crops"
    assert (crop_dir / "page-1-sig_artifact.png").exists()
    assert not list(crop_dir.glob("page-1-signature-*.png"))


def test_table_geometry_uses_source_grid_line_ratios(tmp_path) -> None:
    image_path = tmp_path / "table.png"
    image = Image.new("L", (1000, 700), 255)
    draw = ImageDraw.Draw(image)
    verticals = [40, 160, 440, 490, 960]
    horizontals = [60, 160, 220, 520, 640]
    for x in verticals:
        draw.line((x, horizontals[0], x, horizontals[-1]), fill=0, width=3)
    for y in horizontals:
        draw.line((verticals[0], y, verticals[-1], y), fill=0, width=3)
    image.save(image_path)
    block = LayoutBlock(
        id="table",
        type="table",
        bbox=BoundingBox(x=0.04, y=60 / 700, w=0.92, h=580 / 700),
        confidence=1,
    )
    geometry = _detect_table_grid_geometry(block, image_path, rows=4, cols=4)
    assert geometry is not None
    col_ratios, row_ratios = geometry
    assert [round(value, 2) for value in col_ratios] == [0.13, 0.30, 0.05, 0.51]
    assert [round(value, 2) for value in row_ratios] == [0.17, 0.10, 0.51, 0.21]


def test_bbox_normalization_accepts_pixel_and_percent_coordinates() -> None:
    pixel_box = _normalise_bbox({"x": 300, "y": 600, "w": 900, "h": 120}, image_width=2400, image_height=3600)
    percent_box = _normalise_bbox({"x": 12.5, "y": 16.6, "w": 37.5, "h": 3.3}, image_width=2400, image_height=3600)
    assert pixel_box == {"x": 0.125, "y": 1 / 6, "w": 0.375, "h": 1 / 30}
    assert percent_box["x"] == 0.125
    assert percent_box["w"] == 0.375


def test_compressed_layout_expands_to_page_area() -> None:
    blocks = [
        LayoutBlock(id="a", type="paragraph", bbox=BoundingBox(x=0.05, y=0.02, w=0.2, h=0.03), confidence=1),
        LayoutBlock(id="b", type="table", bbox=BoundingBox(x=0.03, y=0.07, w=0.44, h=0.18), confidence=1),
        LayoutBlock(id="c", type="artifact", bbox=BoundingBox(x=0.14, y=0.26, w=0.06, h=0.02), confidence=1),
        LayoutBlock(id="d", type="paragraph", bbox=BoundingBox(x=0.26, y=0.27, w=0.18, h=0.04), confidence=1),
    ]
    expanded = _expand_compressed_layout(blocks)
    assert expanded[1].bbox.w > 0.8
    assert expanded[3].bbox.y > 0.75


def test_all_direct_matter_demos_render_and_export() -> None:
    for pdf_source in sorted(DEMO_DIRECT_MATTER.glob("*.pdf")):
        job_dir = create_job_dir()
        pdf = save_demo_pdf(job_dir, pdf_source)
        pages = render_pdf_pages(pdf, job_dir / "pages")
        layout = extract_layout(pdf, pages, model="gemini-3-flash-preview", force_demo=True)
        write_layout(job_dir, layout)
        result = export_all(layout, job_dir / "exports", include_bijoy=False)
        assert len(layout.pages) >= 1
        assert Path(result["faithful"]).exists(), pdf_source.name
        assert Path(result["column_ready"]).exists(), pdf_source.name
