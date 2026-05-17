from __future__ import annotations

from pathlib import Path

from docx import Document

from .models import BlockType, BoundingBox, DocumentLayout, LayoutBlock, PageLayout, TableBlock, TableCell
from .storage import DEMO_PREPARE_MATTER


REFERENCE_MATCHES = {
    "02": "02. DSCC-6x4.docx",
    "03": "03. Gopalgonj.docx",
    "04": "04. Railway 7x3.docx",
    "05": "05. Rajshahi Medical College 5x4.docx",
    "06": "06. RHD 10x4.docx",
}


def build_demo_layout(source_pdf: Path, pages: list[PageLayout], reason: str = "missing_api_key") -> DocumentLayout:
    prefix = source_pdf.name[:2]
    reference_name = REFERENCE_MATCHES.get(prefix)
    blocks: list[LayoutBlock] = []
    if reason == "forced_demo":
        warnings = ["Demo/reference mode is on; Gemini OCR was skipped for this conversion."]
        placeholder_text = "Uncheck Demo/reference mode and reconvert this PDF to use Gemini OCR."
    else:
        warnings = ["No Gemini API key found; generated demo layout from local reference DOCX where available."]
        placeholder_text = "Add GEMINI_API_KEY in .env.local or save an API key in the app, then reconvert this PDF."

    if reference_name and (DEMO_PREPARE_MATTER / reference_name).exists():
        blocks = _blocks_from_docx(DEMO_PREPARE_MATTER / reference_name)
    else:
        warnings.append("No matching reference DOCX found; export contains review placeholders.")
        blocks = [
            LayoutBlock(
                id="fallback-heading-1",
                type=BlockType.HEADING,
                text=f"OCR review required: {source_pdf.name}",
                bbox=BoundingBox(x=0.12, y=0.12, w=0.76, h=0.06),
                confidence=0.1,
                language="mixed",
            ),
            LayoutBlock(
                id="fallback-paragraph-1",
                type=BlockType.PARAGRAPH,
                text=placeholder_text,
                bbox=BoundingBox(x=0.12, y=0.22, w=0.76, h=0.08),
                confidence=0.1,
                language="en",
            ),
        ]

    if pages:
        pages[0].blocks = blocks
        for page in pages[1:]:
            page.blocks = []

    return DocumentLayout(
        source_pdf=str(source_pdf),
        pages=pages,
        warnings=warnings,
        ocr_provider="demo-reference",
        ocr_model="local-reference",
        demo_mode=True,
    )


def _blocks_from_docx(path: Path) -> list[LayoutBlock]:
    doc = Document(str(path))
    blocks: list[LayoutBlock] = []
    y = 0.08
    para_index = 0
    table_index = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        block_type = BlockType.HEADING if para_index < 4 or len(text) < 45 and para_index < 8 else BlockType.PARAGRAPH
        blocks.append(
            LayoutBlock(
                id=f"ref-p-{para_index + 1}",
                type=block_type,
                text=text,
                bbox=BoundingBox(x=0.1, y=min(y, 0.92), w=0.8, h=0.035 if block_type == BlockType.HEADING else 0.055),
                confidence=0.85,
                language="mixed",
            )
        )
        para_index += 1
        y += 0.04 if block_type == BlockType.HEADING else 0.065

    for table in doc.tables:
        cells: list[TableCell] = []
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cells.append(TableCell(row=r_idx, col=c_idx, text=cell.text.strip(), confidence=0.85))
        row_count = len(table.rows)
        col_count = max((len(row.cells) for row in table.rows), default=0)
        blocks.append(
            LayoutBlock(
                id=f"ref-table-{table_index + 1}",
                type=BlockType.TABLE,
                bbox=BoundingBox(x=0.08, y=min(y, 0.88), w=0.84, h=min(0.28, 0.04 * max(1, row_count))),
                confidence=0.85,
                language="mixed",
                table=TableBlock(cells=cells, row_count=row_count, col_count=col_count),
            )
        )
        table_index += 1
        y += 0.04 * max(1, row_count) + 0.04

    return blocks
