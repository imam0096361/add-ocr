from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

from .bijoy import split_bijoy_font_runs
from .models import BlockType, BoundingBox, DocumentLayout, LayoutBlock, TableBlock
from .validation import validate_layout


A4_WIDTH_CM = 21.0
A4_HEIGHT_CM = 29.7
UNICODE_BANGLA_FONT = "Noto Serif Bengali"
BIJOY_FONT = "SutonnyMJ"
EN_FONT = "Times New Roman"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


def export_all(layout: DocumentLayout, exports_dir: Path, include_bijoy: bool = True) -> dict[str, str | list[str]]:
    exports_dir.mkdir(parents=True, exist_ok=True)
    warnings = validate_layout(layout)
    faithful = exports_dir / "faithful.docx"
    editable_faithful = exports_dir / "editable-faithful.docx"
    column_ready = exports_dir / "column-ready.docx"
    build_visual_exact_docx(layout, faithful)
    build_docx(layout, editable_faithful, mode="faithful", bijoy=False)
    build_docx(layout, column_ready, mode="column", bijoy=False)
    result: dict[str, str | list[str]] = {
        "faithful": str(faithful),
        "editable_faithful": str(editable_faithful),
        "column_ready": str(column_ready),
        "warnings": warnings,
    }
    if include_bijoy:
        bijoy_path = exports_dir / "column-ready-bijoy.docx"
        bijoy_warnings = build_docx(layout, bijoy_path, mode="column", bijoy=True)
        result["bijoy"] = str(bijoy_path)
        result["bijoy_warnings"] = bijoy_warnings
    return result


def build_docx(layout: DocumentLayout, output_path: Path, mode: str, bijoy: bool = False) -> list[str]:
    doc = Document()
    _setup_styles(doc, bijoy=bijoy)
    warnings: list[str] = []

    for page_index, page in enumerate(layout.pages):
        if page_index:
            doc.add_section(WD_SECTION.NEW_PAGE)
        _setup_section(doc.sections[-1], mode=mode)
        page_blocks = _prepare_faithful_blocks(page.blocks) if mode == "faithful" else page.blocks
        signature_blocks = _find_bottom_signature_blocks(page_blocks) if mode == "faithful" else []
        signature_ids = {block.id for block in signature_blocks}
        for block in sorted(page_blocks, key=lambda b: (b.bbox.y, b.bbox.x)):
            if block.id in signature_ids:
                continue
            if mode == "faithful" and signature_blocks and block.type == BlockType.ARTIFACT and block.bbox.y > 0.68:
                continue
            if block.type == BlockType.TABLE and block.table:
                warnings.extend(_add_table(doc, block, Path(page.image_path), compact=(mode == "faithful"), bijoy=bijoy))
            elif block.type == BlockType.ARTIFACT:
                _add_artifact(doc, block, Path(page.image_path), output_path.parent, page_index, faithful=(mode == "faithful"))
            else:
                warnings.extend(_add_text_block(doc, block, faithful=(mode == "faithful"), bijoy=bijoy))
        if signature_blocks:
            warnings.extend(
                _add_signature_pair(
                    doc,
                    signature_blocks,
                    Path(page.image_path),
                    output_path.parent,
                    page_index,
                    faithful=(mode == "faithful"),
                    bijoy=bijoy,
                )
            )

    doc.save(output_path)
    return sorted(set(warnings))


def build_visual_exact_docx(layout: DocumentLayout, output_path: Path) -> None:
    doc = Document()
    for page_index, page in enumerate(layout.pages):
        if page_index:
            doc.add_section(WD_SECTION.NEW_PAGE)
        section = doc.sections[-1]
        section.page_width = Cm(A4_WIDTH_CM)
        section.page_height = Cm(A4_HEIGHT_CM)
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(0)
        section.right_margin = Cm(0)
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(page.image_path), width=Cm(A4_WIDTH_CM))
    doc.save(output_path)


def render_docx_if_possible(docx_path: Path, output_dir: Path) -> str | None:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    output_dir.mkdir(parents=True, exist_ok=True)
    # Kept intentionally conservative; conversion availability differs on Windows.
    return soffice


def _setup_section(section, mode: str) -> None:
    section.page_width = Cm(A4_WIDTH_CM)
    section.page_height = Cm(A4_HEIGHT_CM)
    margin = 1.2 if mode == "faithful" else 1.6
    section.top_margin = Cm(margin)
    section.bottom_margin = Cm(margin)
    section.left_margin = Cm(margin)
    section.right_margin = Cm(margin)


def _setup_styles(doc: Document, bijoy: bool) -> None:
    font_name = BIJOY_FONT if bijoy else UNICODE_BANGLA_FONT
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(10.5)
    _set_east_asia_font(normal, font_name)

    for style_name, size, bold in [("Title", 14, True), ("Heading 1", 13, True), ("Heading 2", 12, True)]:
        style = doc.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.bold = bold
        _set_east_asia_font(style, font_name)


def _set_east_asia_font(style, font_name: str) -> None:
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _add_page_label(doc: Document, page_number: int, demo_mode: bool) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Page {page_number}")
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(120, 120, 120)
    if demo_mode:
        p.add_run(" - demo OCR").font.size = Pt(7)


def _add_text_block(doc: Document, block: LayoutBlock, faithful: bool, bijoy: bool) -> list[str]:
    style = "Heading 1" if block.type == BlockType.HEADING else "Normal"
    p = doc.add_paragraph(style=style)
    p.alignment = _alignment_from_bbox(block)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0 if faithful else 6)
    font_size = _faithful_font_size(block) if faithful else Pt(11)
    if faithful:
        p.paragraph_format.line_spacing = _faithful_line_spacing(block, font_size)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    else:
        p.paragraph_format.line_spacing = 1.15
    if faithful:
        _apply_faithful_paragraph_geometry(p, block)
    return _add_multiline_run(
        p,
        block.text,
        font_name=BIJOY_FONT if bijoy else UNICODE_BANGLA_FONT,
        size=font_size,
        bold=block.type == BlockType.HEADING,
        underline=block.underline,
        bijoy_mixed=bijoy,
    )


def _prepare_faithful_blocks(blocks: list[LayoutBlock]) -> list[LayoutBlock]:
    prepared = [block.model_copy(deep=True) for block in blocks]
    text_blocks = [
        block
        for block in prepared
        if block.type in {BlockType.PARAGRAPH, BlockType.HEADING}
        and block.bbox.h <= 0.035
        and "\n" not in block.text
    ]
    groups: list[list[LayoutBlock]] = []
    for block in sorted(text_blocks, key=lambda item: item.bbox.y + item.bbox.h / 2):
        center = block.bbox.y + block.bbox.h / 2
        for group in groups:
            group_center = sum(item.bbox.y + item.bbox.h / 2 for item in group) / len(group)
            if abs(center - group_center) <= 0.008:
                group.append(block)
                break
        else:
            groups.append([block])
    for group in groups:
        if len(group) < 2:
            continue
        group.sort(key=lambda item: item.bbox.x)
        if group[-1].bbox.x - group[0].bbox.x < 0.25:
            continue
        shared_y = min(item.bbox.y for item in group)
        for block in group:
            block.bbox.y = shared_y
    return prepared


def _faithful_font_size(block: LayoutBlock):
    line_count = max(1, len(block.text.splitlines()))
    box_height_pt = max(6.0, block.bbox.h * A4_HEIGHT_CM * 28.346)
    size = box_height_pt / max(1.18, line_count * 1.18)
    if block.type == BlockType.HEADING:
        size *= 1.03
    return Pt(max(6.5, min(13.0, size)))


def _faithful_line_spacing(block: LayoutBlock, font_size):
    line_count = max(1, len(block.text.splitlines()))
    box_height_pt = max(font_size.pt * line_count, block.bbox.h * A4_HEIGHT_CM * 28.346)
    line_spacing = box_height_pt / line_count
    return Pt(max(font_size.pt * 1.02, min(font_size.pt * 1.45, line_spacing)))


def _add_multiline_run(
    paragraph,
    text: str,
    font_name: str,
    size,
    bold: bool = False,
    underline: bool = False,
    bijoy_mixed: bool = False,
) -> list[str]:
    warnings: list[str] = []
    lines = text.splitlines() or [text]
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        if bijoy_mixed:
            runs, run_warnings = split_bijoy_font_runs(line)
            warnings.extend(run_warnings)
            for segment, role in runs:
                run_font = BIJOY_FONT if role == "bijoy" else EN_FONT
                _add_formatted_run(paragraph, segment, run_font, size, bold, underline)
        else:
            _add_formatted_run(paragraph, line, font_name, size, bold, underline)
    return sorted(set(warnings))


def _add_formatted_run(paragraph, text: str, font_name: str, size, bold: bool = False, underline: bool = False) -> None:
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    run.underline = underline
    _set_run_fonts(run, font_name)
    _preserve_run_spaces(run)


def _preserve_run_spaces(run) -> None:
    for text_node in run._element.xpath(".//w:t"):
        text_node.set(XML_SPACE, "preserve")


def _set_run_fonts(run, font_name: str) -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    latin_font = font_name if font_name == BIJOY_FONT else EN_FONT
    rfonts.set(qn("w:ascii"), latin_font)
    rfonts.set(qn("w:hAnsi"), latin_font)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _apply_faithful_paragraph_geometry(paragraph, block: LayoutBlock) -> None:
    usable_width_cm = A4_WIDTH_CM - 2.4
    left_cm = max(0, min(usable_width_cm - 1.0, block.bbox.x * usable_width_cm))
    right_edge_cm = max(left_cm + 1.0, min(usable_width_cm, (block.bbox.x + block.bbox.w) * usable_width_cm))
    paragraph.paragraph_format.left_indent = Cm(left_cm)
    paragraph.paragraph_format.right_indent = Cm(max(0, usable_width_cm - right_edge_cm))
    _apply_frame_geometry(paragraph, block)


def _apply_frame_geometry(paragraph, block: LayoutBlock) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    frame_pr = p_pr.find(qn("w:framePr"))
    if frame_pr is None:
        frame_pr = OxmlElement("w:framePr")
        p_pr.insert(0, frame_pr)
    page_width = int(A4_WIDTH_CM * 567)
    page_height = int(A4_HEIGHT_CM * 567)
    frame_pr.set(qn("w:w"), str(max(300, int(block.bbox.w * page_width))))
    frame_pr.set(qn("w:h"), str(max(180, int(block.bbox.h * page_height))))
    frame_pr.set(qn("w:x"), str(max(0, int(block.bbox.x * page_width))))
    frame_pr.set(qn("w:y"), str(max(0, int(block.bbox.y * page_height))))
    frame_pr.set(qn("w:hAnchor"), "page")
    frame_pr.set(qn("w:vAnchor"), "page")
    frame_pr.set(qn("w:wrap"), "none")


def _alignment_from_bbox(block: LayoutBlock):
    if block.alignment == "center":
        return WD_ALIGN_PARAGRAPH.CENTER
    if block.alignment == "right":
        return WD_ALIGN_PARAGRAPH.RIGHT
    if block.alignment == "justify":
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    if block.alignment == "left":
        return WD_ALIGN_PARAGRAPH.LEFT
    center = block.bbox.x + block.bbox.w / 2
    if block.bbox.w < 0.7 and 0.35 < center < 0.65:
        return WD_ALIGN_PARAGRAPH.CENTER
    if block.bbox.x > 0.58:
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


def _add_table(doc: Document, block: LayoutBlock, page_image: Path | None, compact: bool, bijoy: bool) -> list[str]:
    table_block = block.table
    if table_block is None:
        return []
    warnings: list[str] = []
    rows = max(1, table_block.row_count)
    cols = max(1, table_block.col_count)
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT if compact else WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    grid = {(cell.row, cell.col): cell for cell in table_block.cells}
    row_heights_cm = _apply_table_geometry(
        table, block, table_block, cols, rows=rows, compact=compact, page_image=page_image
    )
    for r_idx, row in enumerate(table.rows):
        if compact:
            height_cm = row_heights_cm[r_idx] if row_heights_cm and r_idx < len(row_heights_cm) else _estimated_row_height_cm(
                row_idx=r_idx, rows=rows, block=block, table_block=table_block
            )
            row.height = Cm(height_cm)
            _set_row_height_rule(row, height_cm)
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if compact:
                _set_cell_margins(cell, top=28, start=36, bottom=28, end=36)
            else:
                _set_cell_margins(cell, top=55, start=70, bottom=55, end=70)
            text = grid.get((r_idx, c_idx)).text if (r_idx, c_idx) in grid else ""
            para = cell.paragraphs[0]
            para.alignment = _table_cell_alignment(text)
            warnings.extend(
                _add_multiline_run(
                    para,
                    text,
                    font_name=BIJOY_FONT if bijoy else UNICODE_BANGLA_FONT,
                    size=Pt(7.2 if compact else 9.5),
                    bijoy_mixed=bijoy,
                )
            )
            _set_cell_shading(cell, "F2F2F2" if r_idx == 0 else "FFFFFF")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return sorted(set(warnings))


def _table_cell_alignment(text: str):
    stripped = text.strip()
    if stripped in {":", ""} or len(stripped) < 24:
        return WD_ALIGN_PARAGRAPH.CENTER
    if len(stripped) > 55:
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT


def _apply_table_geometry(
    table,
    block: LayoutBlock,
    table_block: TableBlock | None = None,
    cols: int = 0,
    rows: int = 0,
    compact: bool = False,
    page_image: Path | None = None,
) -> list[float] | None:
    usable_width_cm = A4_WIDTH_CM - 2.4
    detected = _detect_table_grid_geometry(block, page_image, rows=rows, cols=cols) if page_image else None
    width_cm = max(4.0, min(usable_width_cm, block.bbox.w * A4_WIDTH_CM))
    if detected and _grid_geometry_is_plausible(detected[0], table_block, cols):
        col_ratios, row_ratios = detected
        widths_dxa = [max(120, int(width_cm * ratio * 567)) for ratio in col_ratios]
        total_height_cm = max(0.5, min(A4_HEIGHT_CM - 2.4, block.bbox.h * A4_HEIGHT_CM))
        row_heights_cm = [max(0.22, total_height_cm * ratio) for ratio in row_ratios]
    elif _looks_like_colon_form_table(table_block, cols):
        # Source government notice forms usually use a very narrow serial
        # column and a punctuation-only colon column. This is only a fallback
        # when image grid-line detection cannot recover the source geometry.
        width_cm = max(width_cm, usable_width_cm * 0.94)
        weights = [0.045, 0.245, 0.028, 0.682]
        min_widths = [260, 900, 120, 1800]
        widths_dxa = [max(min_widths[index], int(width_cm * weight * 567)) for index, weight in enumerate(weights[:cols])]
        row_heights_cm = None
    else:
        weights = _content_width_weights(table_block, cols)
        min_widths = [180] * cols
        widths_dxa = [max(min_widths[index], int(width_cm * weight * 567)) for index, weight in enumerate(weights[:cols])]
        row_heights_cm = None
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width_dxa = widths_dxa[index] if index < len(widths_dxa) else widths_dxa[-1]
            _set_cell_width(cell, width_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    if compact:
        _apply_floating_table_position(tbl_pr, block)
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width_dxa in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width_dxa))
        tbl_grid.append(grid_col)
    return row_heights_cm


def _grid_geometry_is_plausible(col_ratios: list[float], table_block: TableBlock | None, cols: int) -> bool:
    if len(col_ratios) != cols or not table_block:
        return True
    longest = _longest_text_by_column(table_block, cols)
    for index, ratio in enumerate(col_ratios):
        if _is_punctuation_column(table_block, index):
            continue
        if longest[index] >= 18 and ratio < 0.075:
            return False
        if longest[index] >= 32 and ratio < 0.12:
            return False
    return True


def _looks_like_colon_form_table(table_block: TableBlock | None, cols: int) -> bool:
    if not table_block or cols != 4 or table_block.row_count < 2:
        return False
    colon_cells = [
        cell.text.strip()
        for cell in table_block.cells
        if cell.col == 2 and cell.text.strip()
    ]
    if not colon_cells:
        return False
    colon_like = sum(1 for text in colon_cells if text in {":", "ঃ"})
    return colon_like >= max(1, len(colon_cells) * 0.7)


def _content_width_weights(table_block: TableBlock | None, cols: int) -> list[float]:
    if cols <= 0:
        return []
    if not table_block:
        return [1 / cols] * cols
    longest = _longest_text_by_column(table_block, cols)
    scores: list[float] = []
    for index, length in enumerate(longest):
        if _is_punctuation_column(table_block, index):
            scores.append(1.2)
        elif length <= 4:
            scores.append(4.0)
        elif length <= 12:
            scores.append(7.0)
        else:
            scores.append(float(min(42, max(10, length))))
    total = sum(scores) or cols
    weights = [score / total for score in scores]
    return _enforce_minimum_widths(weights, table_block, cols)


def _enforce_minimum_widths(weights: list[float], table_block: TableBlock, cols: int) -> list[float]:
    minimums: list[float] = []
    longest = _longest_text_by_column(table_block, cols)
    for index, length in enumerate(longest):
        if _is_punctuation_column(table_block, index):
            minimums.append(0.025)
        elif length >= 32:
            minimums.append(0.16)
        elif length >= 18:
            minimums.append(0.10)
        else:
            minimums.append(0.055)
    minimum_total = sum(minimums)
    if minimum_total >= 0.92:
        return [value / minimum_total for value in minimums]
    remaining = 1.0 - minimum_total
    extra_scores = [max(0.0, weight - minimums[index]) for index, weight in enumerate(weights)]
    extra_total = sum(extra_scores)
    if extra_total <= 0:
        extra_scores = [1.0] * cols
        extra_total = float(cols)
    return [minimums[index] + remaining * (extra_scores[index] / extra_total) for index in range(cols)]


def _longest_text_by_column(table_block: TableBlock, cols: int) -> list[int]:
    longest = [0] * cols
    for cell in table_block.cells:
        if 0 <= cell.col < cols:
            length = max((_visual_text_length(line) for line in cell.text.splitlines()), default=0)
            longest[cell.col] = max(longest[cell.col], length)
    return longest


def _visual_text_length(text: str) -> int:
    return len(text.replace(" ", ""))


def _is_punctuation_column(table_block: TableBlock, col_index: int) -> bool:
    texts = [cell.text.strip() for cell in table_block.cells if cell.col == col_index and cell.text.strip()]
    return bool(texts) and sum(1 for text in texts if text in {":", "ঃ"}) >= max(1, len(texts) * 0.7)


def _apply_floating_table_position(tbl_pr, block: LayoutBlock) -> None:
    tblp = tbl_pr.find(qn("w:tblpPr"))
    if tblp is None:
        tblp = OxmlElement("w:tblpPr")
        tbl_pr.insert(0, tblp)
    page_width = int(A4_WIDTH_CM * 567)
    page_height = int(A4_HEIGHT_CM * 567)
    tblp.set(qn("w:leftFromText"), "0")
    tblp.set(qn("w:rightFromText"), "0")
    tblp.set(qn("w:topFromText"), "0")
    tblp.set(qn("w:bottomFromText"), "0")
    tblp.set(qn("w:horzAnchor"), "page")
    tblp.set(qn("w:vertAnchor"), "page")
    tblp.set(qn("w:tblpX"), str(max(0, int(block.bbox.x * page_width))))
    tblp.set(qn("w:tblpY"), str(max(0, int(block.bbox.y * page_height))))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "0")
    overlap = tbl_pr.find(qn("w:tblOverlap"))
    if overlap is None:
        overlap = OxmlElement("w:tblOverlap")
        tbl_pr.append(overlap)
    overlap.set(qn("w:val"), "never")


def _detect_table_grid_geometry(
    block: LayoutBlock, page_image: Path | None, rows: int, cols: int
) -> tuple[list[float], list[float]] | None:
    if not page_image or not page_image.exists() or rows < 1 or cols < 1:
        return None
    try:
        with Image.open(page_image) as img:
            gray = img.convert("L")
            image_width, image_height = gray.size
            pad_x = max(2, int(block.bbox.w * image_width * 0.006))
            pad_y = max(2, int(block.bbox.h * image_height * 0.006))
            left = max(0, int(block.bbox.x * image_width) - pad_x)
            top = max(0, int(block.bbox.y * image_height) - pad_y)
            right = min(image_width, int((block.bbox.x + block.bbox.w) * image_width) + pad_x)
            bottom = min(image_height, int((block.bbox.y + block.bbox.h) * image_height) + pad_y)
            if right <= left or bottom <= top:
                return None
            crop = gray.crop((left, top, right, bottom))
    except OSError:
        return None

    crop_width, crop_height = crop.size
    if crop_width < 40 or crop_height < 40:
        return None
    vertical_runs = _longest_dark_runs(crop, axis="vertical")
    horizontal_runs = _longest_dark_runs(crop, axis="horizontal")
    vertical_groups = _line_groups_from_counts(vertical_runs, threshold=max(12, int(crop_height * 0.14)))
    horizontal_groups = _line_groups_from_counts(horizontal_runs, threshold=max(12, int(crop_width * 0.45)))
    vertical_lines = _choose_strong_grid_lines(vertical_groups, expected=cols + 1, extent=crop_width)
    horizontal_lines = _choose_strong_grid_lines(horizontal_groups, expected=rows + 1, extent=crop_height)
    if len(vertical_lines) != cols + 1 or len(horizontal_lines) != rows + 1:
        return None
    col_widths = [vertical_lines[index + 1] - vertical_lines[index] for index in range(cols)]
    row_heights = [horizontal_lines[index + 1] - horizontal_lines[index] for index in range(rows)]
    if min(col_widths) <= 2 or min(row_heights) <= 2:
        return None
    col_total = sum(col_widths)
    row_total = sum(row_heights)
    return [width / col_total for width in col_widths], [height / row_total for height in row_heights]


def _longest_dark_runs(image, axis: str) -> list[int]:
    width, height = image.size
    px = image.load()
    runs: list[int] = []
    if axis == "vertical":
        for x in range(width):
            best = 0
            current = 0
            for y in range(height):
                if px[x, y] < 150:
                    current += 1
                    best = max(best, current)
                else:
                    current = 0
            runs.append(best)
        return runs
    for y in range(height):
        best = 0
        current = 0
        for x in range(width):
            if px[x, y] < 150:
                current += 1
                best = max(best, current)
            else:
                current = 0
        runs.append(best)
    return runs


def _line_positions_from_counts(counts: list[int], threshold: int) -> list[int]:
    return [position for position, _score in _line_groups_from_counts(counts, threshold)]


def _line_groups_from_counts(counts: list[int], threshold: int) -> list[tuple[int, int]]:
    positions: list[int] = []
    groups: list[tuple[int, int]] = []
    start: int | None = None
    best_index = 0
    best_count = -1
    for index, count in enumerate(counts):
        if count >= threshold:
            if start is None:
                start = index
                best_index = index
                best_count = count
            elif count > best_count:
                best_index = index
                best_count = count
        elif start is not None:
            groups.append((best_index, best_count))
            start = None
            best_count = -1
    if start is not None:
        groups.append((best_index, best_count))
    return groups


def _choose_strong_grid_lines(groups: list[tuple[int, int]], expected: int, extent: int) -> list[int]:
    if expected < 2 or extent <= 0:
        return []
    edge_tolerance = max(8, int(extent * 0.03))
    normalized = [(max(0, min(extent, position)), score) for position, score in groups]
    normalized.sort(key=lambda item: item[0])
    left_edge = min((item for item in normalized if item[0] <= edge_tolerance), default=(0, 0), key=lambda item: item[0])
    right_edge = max(
        (item for item in normalized if item[0] >= extent - edge_tolerance),
        default=(extent, 0),
        key=lambda item: item[0],
    )
    middle = [
        item
        for item in normalized
        if item[0] > edge_tolerance and item[0] < extent - edge_tolerance
    ]
    chosen: list[tuple[int, int]] = [left_edge, right_edge]
    for position, score in sorted(middle, key=lambda item: item[1], reverse=True):
        if len(chosen) >= expected:
            break
        if all(abs(position - existing) > 5 for existing, _ in chosen):
            chosen.append((position, score))
    result = sorted(position for position, _score in chosen)
    if len(result) != expected:
        return []
    result[0] = 0
    result[-1] = extent
    if any(result[index + 1] <= result[index] for index in range(len(result) - 1)):
        return []
    return result


def _choose_grid_lines(lines: list[int], expected: int, extent: int) -> list[int]:
    if expected < 2 or extent <= 0:
        return []
    normalized = sorted(set(max(0, min(extent, line)) for line in lines))
    edge_tolerance = max(6, int(extent * 0.025))
    if not normalized or normalized[0] > edge_tolerance:
        normalized.insert(0, 0)
    else:
        normalized[0] = 0
    if normalized[-1] < extent - edge_tolerance:
        normalized.append(extent)
    else:
        normalized[-1] = extent
    if len(normalized) == expected:
        return normalized
    if len(normalized) < expected:
        return []
    return _best_monotonic_grid_subset(normalized, expected, extent)


def _best_monotonic_grid_subset(lines: list[int], expected: int, extent: int) -> list[int]:
    if len(lines) <= expected:
        return lines
    middle = lines[1:-1]
    needed_middle = expected - 2
    if needed_middle <= 0:
        return [0, extent]
    target_gap = extent / (expected - 1)
    selected: list[int] = []
    last = 0
    for slot in range(1, expected - 1):
        target = int(round(slot * target_gap))
        candidates = [line for line in middle if line > last + 2 and line < extent - 2 and line not in selected]
        if not candidates:
            return []
        chosen = min(candidates, key=lambda line: abs(line - target))
        selected.append(chosen)
        last = chosen
    result = [0, *selected, extent]
    if any(result[index + 1] <= result[index] for index in range(len(result) - 1)):
        return []
    return result


def _set_cell_width(cell, width_dxa: int) -> None:
    cell.width = Cm(width_dxa / 567)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def _set_row_height_rule(row, height_cm: float) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(max(80, int(height_cm * 567))))
    tr_height.set(qn("w:hRule"), "exact")


def _set_cell_margins(cell, top: int, start: int, bottom: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _estimated_row_height_cm(row_idx: int, rows: int, block: LayoutBlock, table_block: TableBlock | None = None) -> float:
    total_cm = max(4.0, min(20.5, block.bbox.h * A4_HEIGHT_CM))
    if rows <= 0:
        return 0.45
    if table_block:
        row_weights = _content_row_height_weights(table_block, rows)
        if row_idx < len(row_weights):
            return max(0.32, total_cm * row_weights[row_idx])
    return max(0.32, total_cm / rows)


def _content_row_height_weights(table_block: TableBlock, rows: int) -> list[float]:
    weights: list[float] = []
    for row_index in range(rows):
        row_cells = [cell for cell in table_block.cells if cell.row == row_index]
        line_count = max((len(cell.text.splitlines()) for cell in row_cells), default=1)
        longest_line = max(
            (_visual_text_length(line) for cell in row_cells for line in cell.text.splitlines()),
            default=0,
        )
        weights.append(max(1.0, line_count * 0.85 + min(longest_line, 80) / 120))
    total = sum(weights) or rows
    return [weight / total for weight in weights]


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _add_artifact(
    doc: Document, block: LayoutBlock, page_image: Path, output_dir: Path, page_index: int, faithful: bool = False
) -> None:
    crop_path = Path(block.image_path) if block.image_path else _crop_artifact(block, page_image, output_dir, page_index)
    p = doc.add_paragraph()
    p.alignment = _alignment_from_bbox(block)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if faithful or block.bbox.y > 0.55:
        _apply_faithful_paragraph_geometry(p, _artifact_frame_block(block, page_image) if faithful else block)
    if crop_path and crop_path.exists():
        width_cm = _artifact_image_width_cm(block, crop_path)
        run = p.add_run()
        run.add_picture(str(crop_path), width=Cm(width_cm))
        return
    label = block.artifact_type or "artifact"
    run = p.add_run(f"[{label} kept as image region]")
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(90, 90, 90)


def _find_bottom_signature_blocks(blocks: list[LayoutBlock]) -> list[LayoutBlock]:
    candidates = [
        block
        for block in blocks
        if block.type in {BlockType.PARAGRAPH, BlockType.HEADING}
        and block.bbox.y > 0.72
        and block.bbox.w > 0.22
        and len(block.text.splitlines()) >= 3
    ]
    if len(candidates) < 2:
        return []
    candidates.sort(key=lambda block: (block.bbox.y, block.bbox.x))
    paired = candidates[:2]
    if abs(paired[0].bbox.y - paired[1].bbox.y) > 0.08:
        return []
    return sorted(paired, key=lambda block: block.bbox.x)


def _add_signature_pair(
    doc: Document,
    blocks: list[LayoutBlock],
    page_image: Path,
    output_dir: Path,
    page_index: int,
    faithful: bool,
    bijoy: bool,
) -> list[str]:
    warnings: list[str] = []
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _remove_table_borders(table)
    usable_width_cm = A4_WIDTH_CM - (2.4 if faithful else 3.2)
    widths = [int(usable_width_cm * 0.49 * 567), int(usable_width_cm * 0.49 * 567)]
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            _set_cell_width(cell, widths[index])
            _set_cell_margins(cell, top=0, start=80, bottom=0, end=80)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    for index, block in enumerate(blocks[:2]):
        cell = table.cell(0, index)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        signature_crop = _crop_signature_above_text(block, page_image, output_dir, page_index)
        if signature_crop:
            run = paragraph.add_run()
            run.add_picture(str(signature_crop), width=Cm(_signature_image_width_cm(block)))
            paragraph.add_run().add_break()
        warnings.extend(
            _add_multiline_run(
                paragraph,
                block.text,
                font_name=BIJOY_FONT if bijoy else UNICODE_BANGLA_FONT,
                size=Pt(9.4 if faithful else 10.5),
                bold=False,
                underline=block.underline,
                bijoy_mixed=bijoy,
            )
        )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return sorted(set(warnings))


def _remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def _signature_image_width_cm(block: LayoutBlock) -> float:
    return max(2.4, min(4.4, block.bbox.w * A4_WIDTH_CM * 0.58))


def _artifact_image_width_cm(block: LayoutBlock, crop_path: Path) -> float:
    try:
        with Image.open(crop_path) as img:
            aspect_ratio = img.width / max(1, img.height)
    except OSError:
        aspect_ratio = 3.0
    base_width = block.bbox.w * A4_WIDTH_CM
    if block.artifact_type in {"signature", "handwriting"}:
        return max(2.6, min(5.8, max(base_width, aspect_ratio * 0.65)))
    return max(1.4, min(5.0, base_width))


def _crop_signature_above_text(
    block: LayoutBlock, page_image: Path, output_dir: Path, page_index: int
) -> Path | None:
    if not page_image.exists():
        return None
    artifact_dir = output_dir / "artifact-crops"
    artifact_dir.mkdir(parents=True, exist_ok=True)
    with Image.open(page_image) as img:
        gray = img.convert("L")
        width, height = gray.size
        left = int((block.bbox.x + block.bbox.w * 0.22) * width)
        right = int((block.bbox.x + block.bbox.w * 0.80) * width)
        top = int(max(0.0, block.bbox.y - 0.085) * height)
        bottom = int(max(block.bbox.y - 0.026, block.bbox.y - 0.035) * height)
        if right <= left or bottom <= top:
            return None
        crop = gray.crop((left, top, right, bottom))
        px = crop.load()
        crop_width, crop_height = crop.size
        row_ink = [sum(1 for x in range(crop_width) if px[x, y] < 145) for y in range(crop_height)]
        ink_rows = [index for index, count in enumerate(row_ink) if count > 5]
        if not ink_rows:
            return None
        row_segments: list[tuple[int, int]] = []
        start = previous = ink_rows[0]
        for row_index in ink_rows[1:]:
            if row_index - previous > 3:
                row_segments.append((start, previous))
                start = row_index
            previous = row_index
        row_segments.append((start, previous))

        signature_cutoff = row_segments[-1][1]
        for current, next_segment in zip(row_segments, row_segments[1:]):
            if next_segment[0] - current[1] >= 12:
                signature_cutoff = current[1]
                break

        dark_pixels: list[tuple[int, int]] = []
        for y in range(crop_height):
            if y > signature_cutoff:
                continue
            for x in range(crop_width):
                if px[x, y] < 145:
                    dark_pixels.append((x, y))
        if len(dark_pixels) < 80:
            return None
        xs = [point[0] for point in dark_pixels]
        ys = [point[1] for point in dark_pixels]
        pad_x = 18
        pad_y = 14
        tight_left = max(0, min(xs) - pad_x)
        tight_top = max(0, min(ys) - pad_y)
        tight_right = min(crop_width, max(xs) + pad_x)
        tight_bottom = min(crop_height, max(ys) + pad_y)
        if tight_right - tight_left < 20 or tight_bottom - tight_top < 12:
            return None
        final_crop = crop.crop((tight_left, tight_top, tight_right, tight_bottom)).convert("RGB")
        out = artifact_dir / f"page-{page_index + 1}-signature-{block.id}.png"
        final_crop.save(out)
        return out


def _crop_artifact(block: LayoutBlock, page_image: Path, output_dir: Path, page_index: int) -> Path | None:
    if not page_image.exists():
        return None
    artifact_dir = output_dir / "artifact-crops"
    artifact_dir.mkdir(parents=True, exist_ok=True)
    with Image.open(page_image) as img:
        width, height = img.size
        left, top, right, bottom = _artifact_crop_box(block, width, height)
        if right <= left or bottom <= top or (right - left) < 12 or (bottom - top) < 12:
            return None
        crop = img.crop((left, top, right, bottom)).convert("RGB")
        out = artifact_dir / f"page-{page_index + 1}-{block.id}.png"
        crop.save(out)
        return out


def _artifact_frame_block(block: LayoutBlock, page_image: Path) -> LayoutBlock:
    if block.artifact_type not in {"signature", "handwriting"} or not page_image.exists():
        return block
    try:
        with Image.open(page_image) as img:
            width, height = img.size
            left, top, right, bottom = _artifact_crop_box(block, width, height)
    except OSError:
        return block
    frame_block = block.model_copy(deep=True)
    frame_block.bbox = BoundingBox(
        x=max(0.0, min(1.0, left / width)),
        y=max(0.0, min(1.0, top / height)),
        w=max(0.01, min(1.0, (right - left) / width)),
        h=max(0.01, min(1.0, (bottom - top) / height)),
    )
    return frame_block


def _artifact_crop_box(block: LayoutBlock, width: int, height: int) -> tuple[int, int, int, int]:
    left = int(block.bbox.x * width)
    top = int(block.bbox.y * height)
    right = int((block.bbox.x + block.bbox.w) * width)
    bottom = int((block.bbox.y + block.bbox.h) * height)
    box_width = max(1, right - left)
    box_height = max(1, bottom - top)
    if block.artifact_type in {"signature", "handwriting"}:
        expand_x = max(int(width * 0.045), int(box_width * 0.55))
        expand_top = max(int(height * 0.09), int(box_height * 2.50))
        expand_bottom = max(int(height * 0.035), int(box_height * 1.45))
        return (
            max(0, left - expand_x),
            max(0, top - expand_top),
            min(width, right + expand_x),
            min(height, bottom + expand_bottom),
        )
    pad_x = max(4, int(box_width * 0.08))
    pad_y = max(4, int(box_height * 0.12))
    return max(0, left - pad_x), max(0, top - pad_y), min(width, right + pad_x), min(height, bottom + pad_y)


def _tighten_visible_ink_crop(image: Image.Image) -> Image.Image:
    gray = image.convert("L")
    px = gray.load()
    dark_pixels: list[tuple[int, int]] = []
    for y in range(gray.height):
        for x in range(gray.width):
            if px[x, y] < 185:
                dark_pixels.append((x, y))
    if len(dark_pixels) < 50:
        return image
    xs = [point[0] for point in dark_pixels]
    ys = [point[1] for point in dark_pixels]
    pad_x = max(12, int((max(xs) - min(xs)) * 0.10))
    pad_y = max(10, int((max(ys) - min(ys)) * 0.18))
    left = max(0, min(xs) - pad_x)
    top = max(0, min(ys) - pad_y)
    right = min(image.width, max(xs) + pad_x)
    bottom = min(image.height, max(ys) + pad_y)
    if right - left < 16 or bottom - top < 10:
        return image
    return image.crop((left, top, right, bottom))
